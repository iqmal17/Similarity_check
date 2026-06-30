# SimCheck - exporter.py
import os
import pandas as pd
from datetime import datetime

def export_excel(pairs, matrix, doc_names, output_path):
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Ringkasan pasangan
        rows = []
        for p in pairs:
            rows.append({
                'Dokumen A': p['doc_a'],
                'Dokumen B': p['doc_b'],
                'Kemiripan (%)': p['similarity'],
                'Level': p['level']['label'],
                'Jumlah Fragmen Mirip': len(p.get('fragments', []))
            })
        df_pairs = pd.DataFrame(rows)
        df_pairs.to_excel(writer, sheet_name='Ringkasan', index=False)

        # Matrix
        df_matrix = pd.DataFrame(matrix, index=doc_names, columns=doc_names)
        df_matrix = (df_matrix * 100).round(2)
        df_matrix.to_excel(writer, sheet_name='Heatmap Matrix')

        # Detail fragmen
        detail_rows = []
        for p in pairs:
            for f in p.get('fragments', []):
                detail_rows.append({
                    'Doc A': p['doc_a'],
                    'Doc B': p['doc_b'],
                    'Skor %': f['skor'],
                    'Kalimat A': f['kalimat_a'],
                    'Kalimat B': f['kalimat_b']
                })
        if detail_rows:
            pd.DataFrame(detail_rows).to_excel(writer, sheet_name='Fragmen Mirip', index=False)
    return output_path

def export_word(pairs, matrix, doc_names, mode, output_path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    h = doc.add_heading('Laporan SimCheck - Deteksi Kemiripan Tugas', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Tanggal Analisis: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph(f"Mode Deteksi: {mode.capitalize()}")
    doc.add_paragraph(f"Jumlah Dokumen: {len(doc_names)}")
    doc.add_paragraph(f"Total Pasangan Dibandingkan: {len(pairs)}")
    doc.add_paragraph("")

    # Ringkasan tertinggi
    doc.add_heading('Ringkasan Kemiripan Tertinggi', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Dokumen A'
    hdr[1].text = 'Dokumen B'
    hdr[2].text = 'Kemiripan'
    hdr[3].text = 'Level'
    for p in pairs[:15]:
        row = table.add_row().cells
        row[0].text = p['doc_a']
        row[1].text = p['doc_b']
        row[2].text = f"{p['similarity']}%"
        row[3].text = p['level']['label']
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_paragraph("")

    # Heatmap matrix sebagai tabel
    doc.add_heading('Heatmap Matrix (%)', level=1)
    m_table = doc.add_table(rows=len(doc_names)+1, cols=len(doc_names)+1)
    m_table.style = 'Table Grid'
    m_table.cell(0,0).text = ''
    for j, name in enumerate(doc_names):
        m_table.cell(0, j+1).text = name[:18]
    for i, name in enumerate(doc_names):
        m_table.cell(i+1, 0).text = name[:18]
        for j in range(len(doc_names)):
            val = round(matrix[i][j]*100,1)
            m_table.cell(i+1, j+1).text = f"{val}%"
    doc.add_paragraph("")

    # Detail fragmen
    doc.add_heading('Detail Bagian yang Mirip', level=1)
    for p in pairs[:8]:  # batasi 8 pasangan teratas agar file tidak terlalu besar
        if not p.get('fragments'):
            continue
        doc.add_heading(f"{p['doc_a']}  ↔  {p['doc_b']}  —  {p['similarity']}%  ({p['level']['label']})", level=2)
        for idx, f in enumerate(p['fragments'][:10], 1):
            pf = doc.add_paragraph()
            pf.add_run(f"[{idx}] Skor {f['skor']}% \n").bold = True
            pf.add_run("A: ").bold = True
            pf.add_run(f"{f['kalimat_a']}\n")
            pf.add_run("B: ").bold = True
            pf.add_run(f"{f['kalimat_b']}")
        doc.add_paragraph("")
    
    doc.add_paragraph("")
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("— SimCheck • Plagiarism & Similarity Checker —\nwww.simcheck.id")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120,120,120)

    doc.save(output_path)
    return output_path

def export_pdf(pairs, matrix, doc_names, mode, output_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm, leftMargin=1.8*cm, rightMargin=1.8*cm)
    styles = getSampleStyleSheet()
    story = []
    title_style = ParagraphStyle('t', parent=styles['Title'], fontSize=18, textColor=colors.HexColor('#7c3aed'))
    story.append(Paragraph("SimCheck — Laporan Kemiripan Tugas", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Mode: {mode.capitalize()} | Dokumen: {len(doc_names)} | {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    story.append(Spacer(1, 18))

    story.append(Paragraph("Ringkasan Pasangan", styles['Heading2']))
    data = [['Dokumen A','Dokumen B','%','Level']]
    for p in pairs[:20]:
        data.append([p['doc_a'][:28], p['doc_b'][:28], f"{p['similarity']}%", p['level']['label']])
    t = Table(data, colWidths=[150,150,50,80])
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#7c3aed')),
        ('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('GRID',(0,0),(-1,-1),0.5,colors.grey),
        ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, colors.HexColor('#f5f3ff')])
    ]))
    story.append(t)
    story.append(Spacer(1, 18))

    story.append(Paragraph("Heatmap Matrix (%)", styles['Heading2']))
    hm_data = [[''] + [n[:12] for n in doc_names]]
    for i, n in enumerate(doc_names):
        row = [n[:12]] + [f"{round(matrix[i][j]*100,1)}%" for j in range(len(doc_names))]
        hm_data.append(row)
    hm = Table(hm_data)
    hm.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.4,colors.grey),
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#ede9fe')),
        ('BACKGROUND',(0,0),(0,-1),colors.HexColor('#ede9fe')),
    ]))
    story.append(hm)
    story.append(Spacer(1,24))
    story.append(Paragraph("SimCheck • Plagiarism & Similarity Checker", ParagraphStyle('f', parent=styles['Normal'], textColor=colors.grey, alignment=1)))
    doc.build(story)
    return output_path
